#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ohdoc：NAPI 测试用例 Word 文档生成（依赖 python-docx）。

CLI：python ohdoc.py napi-test-doc [选项] / python ohdoc.py csv-test-doc [选项]
  --test-file <*.test.ts>     单文件，生成一篇 docx。
  --test-dir  <目录>          目录内全部 *.test.ts 合并为**一篇** docx（与 --test-file 二选一）。
  --out <路径>                输出文件；省略则 {技能目录}/napitest{YYYYMMDD}.docx。
  --out-dir <目录>            仅当未写 --out 时，将上述默认文件名写到该目录（默认同技能目录）。
  --template / --executor / --exec-date  见 ohdoc/SKILL.md。

核心逻辑：
  - build_napi_test_doc_from_sources(sources, ...)：单文件传 [path] 即可。
  - 解析 suite('…')、跳过 assert/vscode 后的首个 import 路径、test('…', ()=>{…}) 与 // 注释；
    预置条件/步骤识别 parsec.* 与 parsets.*。
  - 版式：文首元数据 + 模板说明一次；每条用「任务n. xxx测试用例」段落，p._p.addnext(tbl)
    使表格紧贴标题；表与表之间 spacer 段（零宽空格 + 段间距）。
  - 多源合并：文首列出全部源路径；每个源文件首次出现前插入「【文件名】」；任务编号全局连续。
  - 无 test 的 .test.ts 打印警告并跳过；全部无有效用例则报错。
  - csv-test-doc：从测试用例列表 CSV + 同一模板生成 docx（默认输出 {技能目录}/csvtest{YYYYMMDD}.docx）。
    填表：用例名称←CSV 用例名称；用例标识←用例编号；所属模块/功能←功能模块；所属攻关任务 固定「投屏分享」；
    用例追溯←功能模块；需求描述←用例名称；预置/步骤/预期←前置条件、步骤描述、预期结果；维护人、执行日期等见 SKILL.md。
    _fill_one_table 使用 data['identifier'] 作为行 0 右栏（缺省同 name，TS 模式不变）。

详述与表格字段映射见同目录 SKILL.md。Word 模板与 CSV 须放在 template/ 下本地自备（根 .gitignore 忽略 *.docx/*.csv，不入库）。
"""

from __future__ import annotations

import argparse
import csv
import re
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.table import Table
except ImportError:
    print("请先安装依赖: pip install -r requirements.txt", file=sys.stderr)
    sys.exit(1)


def _template_path(skill_dir: Path) -> Path:
    tdir = skill_dir / "template"
    candidates = list(tdir.glob("*.docx"))
    if not candidates:
        raise FileNotFoundError(f"未找到模板 docx: {tdir}")
    # 选用「NAPI测试用例模板」（通常比设计文档模板更大）
    return max(candidates, key=lambda p: p.stat().st_size)


def _extract_suite_and_import(ts: str) -> tuple[str, str]:
    suite_m = re.search(r"suite\s*\(\s*['\"]([^'\"]+)['\"]", ts)
    suite_name = suite_m.group(1).strip() if suite_m else "Unknown_Suite"
    imp_path = ""
    for m in re.finditer(
        r"import\s+\*\s+as\s+\w+\s+from\s+['\"]([^'\"]+)['\"]",
        ts,
    ):
        mod = m.group(1).strip()
        if mod in ("assert", "vscode"):
            continue
        imp_path = mod
        break
    return suite_name, imp_path


def _find_test_blocks(ts: str) -> list[tuple[str, str, int]]:
    """
    返回 [(comment_or_empty, test_name, start_index_in_ts), ...]
    comment 为紧邻 test 之前的 // 行（可多行取最后一行非空说明）或空字符串。
    """
    results: list[tuple[str, str, int]] = []
    for m in re.finditer(
        r"(?:^|\n)([\t ]*//[^\n]*\n)*[\t ]*test\s*\(\s*['\"]([^'\"]+)['\"]",
        ts,
    ):
        block_start = m.start()
        prefix = m.group(0)
        name = m.group(2)
        # 取前缀中的 // 注释，合并多行
        comment_lines = re.findall(r"//\s*(.*)", prefix)
        comment = ""
        for line in comment_lines:
            t = line.strip()
            if t:
                comment = t
        results.append((comment, name, block_start))
    return results


def _brace_body(ts: str, start_at: int) -> str:
    """从 test(..., () => { 起，取出箭头函数体（含内部大括号匹配）。"""
    sub = ts[start_at:]
    m = re.search(r"test\s*\(\s*['\"]([^'\"]+)['\"]\s*,\s*\(\)\s*=>\s*\{", sub)
    if not m:
        return ""
    open_brace = start_at + m.end() - 1
    assert ts[open_brace] == "{"
    depth = 0
    i = open_brace
    while i < len(ts):
        c = ts[i]
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                return ts[open_brace + 1 : i]
        i += 1
    return ""


def _guess_module_api(body: str) -> str:
    m = re.search(r"(?:parsec|parsets)\.(\w+)\s*\(", body)
    return m.group(1) if m else "接口"


def _precondition_snippet(body: str) -> str:
    """取第一次 parsec./parsets. 调用之前的 let/const 声明（含模板字符串）。"""
    m_call = re.search(r"(?:parsec|parsets)\.\w+\s*\(", body)
    if not m_call:
        return body.strip()[:4000]
    before = body[: m_call.start()]
    lines = before.splitlines()
    out: list[str] = []
    for line in lines:
        s = line.strip()
        if re.match(r"(let|const|var)\s+", s):
            out.append(line)
        elif out and (line.startswith(" ") or line.startswith("\t") or not s):
            if out:
                out.append(line)
        elif not s:
            if out:
                out.append(line)
    text = "\n".join(out).strip()
    if len(text) > 4500:
        text = text[:4500] + "\n…（截断）"
    return text or "（无单独变量声明，参见用例正文）"


def _input_steps(body: str) -> str:
    calls = re.findall(r"(?:parsec|parsets)\.\w+\s*\([^;]*\)", body, re.DOTALL)
    if not calls:
        return "（调用见代码）"
    seen: list[str] = []
    for c in calls:
        one = " ".join(c.split())
        if one not in seen:
            seen.append(one)
    return "接口调用：\n" + "\n".join(f"{i+1}. {s}" for i, s in enumerate(seen))


def _expected_asserts(body: str) -> str:
    lines = []
    for line in body.splitlines():
        s = line.strip()
        if s.startswith("assert.") or s.startswith("// assert."):
            lines.append(line.strip())
    text = "\n".join(lines)
    if len(text) > 5500:
        text = text[:5500] + "\n…（截断）"
    return text or "（无 assert，参见代码）"


def _spacer_paragraph_after_table():
    """
    插在表格后的空段（带段前/段后间距），避免 Word 中相邻表格视觉粘连；
    使用零宽空格保证段落非完全空，防止被合并。
    """
    return parse_xml(
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:pPr><w:spacing w:before=\"160\" w:after=\"280\"/></w:pPr>"
        '<w:r><w:t xml:space="preserve">\u200b</w:t></w:r>'
        "</w:p>"
    )


def _fill_one_table(table, data: dict[str, str]) -> None:
    """按模板表格行结构填充（7 列）。行 0 右栏为 data['identifier']（缺省同 name）；第 0、1、12 行右栏值写在列 6。"""
    def set_row_merge_text(row_idx: int, pairs: list[tuple[int, str]]) -> None:
        row = table.rows[row_idx]
        for col_idx, val in pairs:
            if col_idx < len(row.cells):
                row.cells[col_idx].text = val

    name = data["name"]
    # 用例标识（模板行 0 右栏）：TS 模式与「用例名称」相同；CSV 模式为「用例编号」
    identifier = data.get("identifier", name)
    suite = data["suite"]
    task = data["task"]
    trace = data["trace"]
    req = data["requirement"]
    desc = data["description"]
    pre = data["precondition"]
    steps = data["steps"]
    expect = data["expected"]
    pass_rule = data["pass_rule"]
    actual = data["actual"]
    executor = data["executor"]
    exec_date = data["exec_date"]
    conclusion = data["conclusion"]
    anomaly = data["anomaly"]

    # 模板行 0/1：列 0=左标签，列 1–3=左值区，列 4–5=右标签（合并），列 6=右值；勿写列 5 以免覆盖标签
    set_row_merge_text(0, [(1, name), (6, identifier)])
    set_row_merge_text(1, [(1, suite), (6, task)])
    # 行 4–11：左标签列保持，右侧大区域多列合并重复——统一写入列 1
    set_row_merge_text(4, [(1, trace)])
    set_row_merge_text(5, [(1, req)])
    set_row_merge_text(6, [(1, desc)])
    set_row_merge_text(7, [(1, pre)])
    set_row_merge_text(8, [(1, steps)])
    set_row_merge_text(9, [(1, expect)])
    set_row_merge_text(10, [(1, pass_rule)])
    set_row_merge_text(11, [(1, actual)])
    # 行 12：列 3–4 为「执行日期」标签，日期写在列 6
    set_row_merge_text(12, [(1, executor), (6, exec_date)])
    set_row_merge_text(13, [(1, conclusion)])
    set_row_merge_text(14, [(1, anomaly)])


def build_napi_test_doc_from_sources(
    sources: list[Path],
    out_docx: Path,
    template: Path | None = None,
    executor: str = "胡瑞涛",
    exec_date: str | None = None,
) -> None:
    """
    由一个或多个 *.test.ts 生成单个 docx。

    多源时：文首列全部路径；每个文件首测前插入「【文件名】」；任务序号 1..N 全局连续。
    单源时：文首仅一条「源文件：…」。
    版式与 SKILL.md 一致（标题段 addnext 表格、表间 spacer）。
    """
    skill_dir = Path(__file__).resolve().parent
    tpl = template or _template_path(skill_dir)
    if not tpl.is_file():
        raise FileNotFoundError(tpl)
    if not sources:
        raise ValueError("sources 为空")

    exec_date = exec_date or datetime.now().strftime("%Y/%m/%d")

    # (源路径, ts全文, suite, import追溯, (comment, test名, start_idx))
    entries: list[tuple[Path, str, str, str, tuple[str, str, int]]] = []
    for test_ts in sources:
        ts_text = test_ts.read_text(encoding="utf-8")
        suite_name, import_path = _extract_suite_and_import(ts_text)
        blocks = _find_test_blocks(ts_text)
        if not blocks:
            print(f"⚠ 跳过（无 test）：{test_ts.name}", file=sys.stderr)
            continue
        for b in blocks:
            entries.append((test_ts, ts_text, suite_name, import_path, b))

    if not entries:
        raise ValueError("未在任何 *.test.ts 中解析到 test(...)")

    base = Document(str(tpl))
    if not base.tables:
        raise ValueError("模板中无表格，请检查 任务7.2：NAPI测试用例模板.docx")

    template_table = base.tables[0]
    note_text = ""
    if len(base.paragraphs) > 1:
        note_text = base.paragraphs[1].text.strip()

    out = Document()
    out.add_paragraph("NAPI 测试用例（由 ohdoc 自 TypeScript 测试生成）")
    if len(sources) == 1:
        out.add_paragraph(f"源文件：{sources[0].as_posix()}")
    else:
        out.add_paragraph("源文件（多文件合并为一篇，任务编号连续）：")
        for p in sources:
            out.add_paragraph(f"  • {p.as_posix()}")
    out.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    out.add_paragraph("")
    if note_text:
        out.add_paragraph(note_text)
        out.add_paragraph("")

    n_total = len(entries)
    prev_ts: Path | None = None
    for global_i, (test_ts, ts_text, suite_name, import_path, (comment, test_name, start_idx)) in enumerate(
        entries
    ):
        if test_ts != prev_ts:
            if prev_ts is not None:
                out.add_paragraph("")
            out.add_paragraph(f"【{test_ts.name}】")
            prev_ts = test_ts

        task_no = global_i + 1
        title = f"任务{task_no}. {test_name}测试用例"
        p_title = out.add_paragraph(title)

        body = _brace_body(ts_text, start_idx)
        api = _guess_module_api(body)

        data = {
            "name": test_name,
            "suite": suite_name,
            "task": "NAPI",
            "trace": import_path,
            "requirement": comment or f"验证 {api} 行为",
            "description": f"本测试用例主要对 {api} 进行验证（{test_name}）。",
            "precondition": _precondition_snippet(body),
            "steps": _input_steps(body),
            "expected": _expected_asserts(body),
            "pass_rule": "判断正确：各 assert 与预期一致",
            "actual": "见后图",
            "executor": executor,
            "exec_date": exec_date,
            "conclusion": "通过",
            "anomaly": "无",
        }

        tbl_el = deepcopy(template_table._tbl)
        p_title._p.addnext(tbl_el)
        new_table = Table(tbl_el, out._body)
        _fill_one_table(new_table, data)
        if global_i < n_total - 1:
            tbl_el.addnext(_spacer_paragraph_after_table())

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    out.save(str(out_docx))


def build_napi_test_doc(
    test_ts: Path,
    out_docx: Path,
    template: Path | None = None,
    executor: str = "胡瑞涛",
    exec_date: str | None = None,
) -> None:
    build_napi_test_doc_from_sources(
        [test_ts], out_docx, template=template, executor=executor, exec_date=exec_date
    )


def _norm_csv_header_key(key: str) -> str:
    s = (key or "").replace("\ufeff", "").strip().rstrip("*").strip()
    return s


def _csv_row_to_dict(raw: dict[str, str | None]) -> dict[str, str]:
    out: dict[str, str] = {}
    for k, v in raw.items():
        nk = _norm_csv_header_key(str(k))
        if not nk:
            continue
        out[nk] = (v or "").strip() if isinstance(v, str) else ""
    return out


def _load_csv_testcase_rows(csv_path: Path) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    with csv_path.open(encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for raw in reader:
            row = _csv_row_to_dict(raw)
            if not row.get("用例名称", "").strip():
                continue
            rows.append(row)
    return rows


def build_csv_test_doc(
    csv_path: Path,
    out_docx: Path,
    template: Path | None = None,
    exec_date: str | None = None,
) -> None:
    """
    从「测试用例列表」CSV（列含 用例名称、功能模块、用例编号、维护人、前置条件、步骤描述、预期结果 等）
    与 NAPI 测试用例模板生成单篇 docx。标题「任务n. {用例名称}测试用例」。
    填表：用例标识=用例编号；所属攻关任务=「投屏分享」；用例追溯=功能模块；其余见 SKILL.md。
    """
    skill_dir = Path(__file__).resolve().parent
    tpl = template or _template_path(skill_dir)
    if not tpl.is_file():
        raise FileNotFoundError(tpl)
    exec_date = exec_date or "2026/04/10"

    cases = _load_csv_testcase_rows(csv_path)
    if not cases:
        raise ValueError("CSV 中无有效用例行（用例名称为空或文件无数据）")

    base = Document(str(tpl))
    if not base.tables:
        raise ValueError("模板中无表格，请检查 NAPI测试用例模板.docx")

    template_table = base.tables[0]
    note_text = ""
    if len(base.paragraphs) > 1:
        note_text = base.paragraphs[1].text.strip()

    out = Document()
    out.add_paragraph("NAPI 测试用例（由 ohdoc 自 CSV 列表生成）")
    out.add_paragraph(f"源 CSV：{csv_path.resolve().as_posix()}")
    out.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    out.add_paragraph("")
    if note_text:
        out.add_paragraph(note_text)
        out.add_paragraph("")

    n_total = len(cases)
    for i, row in enumerate(cases):
        case_name = row.get("用例名称", "").strip()
        module = row.get("功能模块", "").strip()
        case_id = row.get("用例编号", "").strip()
        maintainer = row.get("维护人", "").strip()
        pre = row.get("前置条件", "").strip()
        steps = row.get("步骤描述", "").strip()
        expect = row.get("预期结果", "").strip()

        task_no = i + 1
        title = f"任务{task_no}. {case_name}测试用例"
        p_title = out.add_paragraph(title)

        data = {
            "name": case_name,
            "identifier": case_id or "—",
            "suite": module or "—",
            "task": "投屏分享",
            "trace": module or "—",
            "requirement": case_name,
            "description": f"本测试用例验证「{case_name}」。",
            "precondition": pre or "（无）",
            "steps": steps or "（无）",
            "expected": expect or "（无）",
            "pass_rule": "符合预期结果",
            "actual": "见后图",
            "executor": maintainer or "—",
            "exec_date": exec_date,
            "conclusion": "通过",
            "anomaly": "无",
        }

        tbl_el = deepcopy(template_table._tbl)
        p_title._p.addnext(tbl_el)
        new_table = Table(tbl_el, out._body)
        _fill_one_table(new_table, data)
        if i < n_total - 1:
            tbl_el.addnext(_spacer_paragraph_after_table())

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    out.save(str(out_docx))


def main() -> None:
    parser = argparse.ArgumentParser(description="ohdoc：NAPI 测试用例文档生成")
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_nt = sub.add_parser(
        "napi-test-doc",
        help="从 TS 测试文件与模板生成 napitest+日期.docx",
    )
    p_nt.add_argument(
        "--test-file",
        default="",
        help="单个 TypeScript 测试文件路径（如 parsec.test.ts）；与 --test-dir 二选一",
    )
    p_nt.add_argument(
        "--test-dir",
        default="",
        help="目录：合并该目录下全部 *.test.ts 为**一个** docx（与 --test-file 二选一）",
    )
    p_nt.add_argument(
        "--out",
        default="",
        help="单文件模式：输出 docx 路径；默认本技能目录下 napitestYYYYMMDD.docx",
    )
    p_nt.add_argument(
        "--out-dir",
        default="",
        help="目录模式：未指定 --out 时，将 napitest{日期}.docx 写到该目录（默认 ohdoc 技能目录）",
    )
    p_nt.add_argument(
        "--template",
        default="",
        help="覆盖默认「NAPI测试用例模板.docx」",
    )
    p_nt.add_argument("--executor", default="胡瑞涛", help="执行人员")
    p_nt.add_argument("--exec-date", default="", help="执行日期，默认今天 YYYY/MM/DD")

    p_csv = sub.add_parser(
        "csv-test-doc",
        help="从测试用例列表 CSV 与模板生成 docx（字段映射见 SKILL.md）",
    )
    p_csv.add_argument(
        "--csv",
        dest="csv_path",
        required=True,
        help="测试用例列表 CSV（UTF-8，含表头；列含 用例名称、功能模块、用例编号、维护人、前置条件、步骤描述、预期结果 等）",
    )
    p_csv.add_argument(
        "--out",
        default="",
        help="输出 docx；默认本技能目录下 csvtestYYYYMMDD.docx",
    )
    p_csv.add_argument(
        "--template",
        default="",
        help="覆盖默认「NAPI测试用例模板.docx」",
    )
    p_csv.add_argument(
        "--exec-date",
        default="2026/04/10",
        help="表格「执行日期」，默认 2026/04/10（YYYY/MM/DD）",
    )

    args = parser.parse_args()
    if args.cmd == "csv-test-doc":
        skill_dir = Path(__file__).resolve().parent
        tpl = Path(args.template).resolve() if args.template else None
        csv_file = Path(args.csv_path).resolve()
        if not csv_file.is_file():
            print(f"❌ 文件不存在: {csv_file}", file=sys.stderr)
            sys.exit(1)
        day = datetime.now().strftime("%Y%m%d")
        if args.out:
            out_path = Path(args.out).resolve()
        else:
            out_path = skill_dir / f"csvtest{day}.docx"
        try:
            build_csv_test_doc(
                csv_file,
                out_path,
                template=tpl,
                exec_date=args.exec_date.strip() or "2026/04/10",
            )
            print(f"已生成: {out_path}")
        except Exception as e:
            print(f"❌ 生成失败: {e}", file=sys.stderr)
            sys.exit(1)
    elif args.cmd == "napi-test-doc":
        skill_dir = Path(__file__).resolve().parent
        tpl = Path(args.template).resolve() if args.template else None
        exec_d = args.exec_date or None

        test_dir = (args.test_dir or "").strip()
        test_file = (args.test_file or "").strip()
        if bool(test_dir) == bool(test_file):
            p_nt.error("请只指定其一：--test-file <单文件> 或 --test-dir <目录>")

        if test_dir:
            d = Path(test_dir).resolve()
            if not d.is_dir():
                print(f"❌ 不是目录: {d}", file=sys.stderr)
                sys.exit(1)
            files = sorted(d.glob("*.test.ts"))
            if not files:
                print(f"❌ 未找到 *.test.ts: {d}", file=sys.stderr)
                sys.exit(1)
            day = datetime.now().strftime("%Y%m%d")
            out_dir = Path(args.out_dir).resolve() if args.out_dir else skill_dir
            out_dir.mkdir(parents=True, exist_ok=True)
            if args.out:
                out_path = Path(args.out).resolve()
            else:
                out_path = out_dir / f"napitest{day}.docx"
            try:
                build_napi_test_doc_from_sources(
                    files,
                    out_path,
                    template=tpl,
                    executor=args.executor,
                    exec_date=exec_d,
                )
                print(f"已生成（合并 {len(files)} 个源文件）: {out_path}")
            except Exception as e:
                print(f"❌ 生成失败: {e}", file=sys.stderr)
                sys.exit(1)
        else:
            tf = Path(test_file).resolve()
            if args.out:
                out_path = Path(args.out).resolve()
            else:
                out_path = skill_dir / f"napitest{datetime.now().strftime('%Y%m%d')}.docx"
            if args.out_dir:
                print("⚠ 单文件模式下忽略 --out-dir", file=sys.stderr)
            build_napi_test_doc(
                tf,
                out_path,
                template=tpl,
                executor=args.executor,
                exec_date=exec_d,
            )
            print(f"已生成: {out_path}")


if __name__ == "__main__":
    main()
